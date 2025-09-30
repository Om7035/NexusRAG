from typing import List
from .base import BaseParser, Document


class WordParser(BaseParser):
    """Word document parser using the python-docx library."""
    
    def parse(self, file_path: str) -> List[Document]:
        """Parse a Word document and return a list of Document objects.
        
        Args:
            file_path (str): Path to the Word file to parse
            
        Returns:
            List[Document]: List of parsed documents
        """
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError(
                "To use WordParser, you need to install the python-docx library. "
                "Please run: pip install python-docx"
            )
        
        # Open the Word document
        doc = DocxDocument(file_path)
        
        # Extract paragraphs
        paragraphs = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # Only add non-empty paragraphs
                paragraphs.append(paragraph.text)
        
        # Extract tables
        tables = []
        for table in doc.tables:
            table_content = ""
            for row in table.rows:
                row_content = "\t".join(cell.text for cell in row.cells)
                table_content += row_content + "\n"
            if table_content.strip():
                tables.append(table_content)
        
        # Combine all content
        all_content = paragraphs + tables
        
        # Convert to Document objects
        documents = []
        for i, content in enumerate(all_content):
            metadata = {
                "source": file_path,
                "content_type": "table" if i >= len(paragraphs) else "paragraph",
                "element_index": i
            }
            document = Document(content=content, metadata=metadata)
            documents.append(document)
            
        return documents
